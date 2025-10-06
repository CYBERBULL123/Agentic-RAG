"""
Advanced Document Processing with ETL Pipeline
Handles PDF, DOCX, TXT, CSV, XLSX, HTML, MD files with cleaning and preprocessing
"""

import os
import io
import re
import json
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
import hashlib

# Document processing imports
import PyPDF2
from docx import Document
import pandas as pd
import html2text
from bs4 import BeautifulSoup
from loguru import logger

# Text processing imports
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer

# Download required NLTK data (first time only)
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('wordnet')


class DocumentProcessor:
    """Advanced document processor with ETL capabilities."""
    
    def __init__(self, 
                 chunk_size: int = 1000,
                 chunk_overlap: int = 200,
                 min_chunk_size: int = 100,
                 clean_text: bool = True):
        """
        Initialize document processor.
        
        Args:
            chunk_size: Maximum size of each text chunk
            chunk_overlap: Number of characters to overlap between chunks
            min_chunk_size: Minimum size for a valid chunk
            clean_text: Whether to perform text cleaning
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        self.clean_text = clean_text
        
        # Initialize text processing tools
        if clean_text:
            self.lemmatizer = WordNetLemmatizer()
            self.stop_words = set(stopwords.words('english'))
        
        # Supported file extensions and their processors
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_txt,
            '.csv': self._process_csv,
            '.xlsx': self._process_xlsx,
            '.html': self._process_html,
            '.md': self._process_markdown
        }
        
        logger.info(f"Initialized DocumentProcessor with chunk_size={chunk_size}")
    
    def process_uploaded_file(self, 
                            file_content: bytes, 
                            filename: str,
                            extract_metadata: bool = True) -> Dict[str, Any]:
        """
        Process an uploaded file through the ETL pipeline.
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename
            extract_metadata: Whether to extract document metadata
            
        Returns:
            Processing result with extracted text, chunks, and metadata
        """
        try:
            logger.info(f"Processing file: {filename}")
            
            # Get file extension
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext not in self.supported_extensions:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {file_ext}",
                    "supported_types": list(self.supported_extensions.keys())
                }
            
            # Extract text using appropriate processor
            processor = self.supported_extensions[file_ext]
            raw_text = processor(file_content)
            
            if not raw_text or len(raw_text.strip()) < 10:
                return {
                    "success": False,
                    "error": "No readable text found in document"
                }
            
            # Generate document hash for deduplication
            doc_hash = self._generate_document_hash(raw_text)
            
            # Clean and preprocess text
            if self.clean_text:
                cleaned_text = self._clean_text(raw_text)
            else:
                cleaned_text = raw_text
            
            # Create text chunks
            chunks = self._create_chunks(cleaned_text)
            
            # Extract metadata if requested
            metadata = {}
            if extract_metadata:
                metadata = self._extract_metadata(raw_text, filename, file_ext)
            
            # Prepare result
            result = {
                "success": True,
                "filename": filename,
                "file_type": file_ext,
                "document_hash": doc_hash,
                "raw_text": raw_text,
                "cleaned_text": cleaned_text,
                "chunks": chunks,
                "metadata": metadata,
                "processing_stats": {
                    "original_length": len(raw_text),
                    "cleaned_length": len(cleaned_text),
                    "num_chunks": len(chunks),
                    "avg_chunk_size": sum(len(chunk) for chunk in chunks) // len(chunks) if chunks else 0
                },
                "processed_at": datetime.now().isoformat()
            }
            
            logger.info(f"Successfully processed {filename}: {len(chunks)} chunks created")
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename
            }
    
    def _process_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {e}")
                    continue
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return ""
    
    def _process_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append("[Table]\n" + "\n".join(table_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            return ""
    
    def _process_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use errors='ignore'
            return file_content.decode('utf-8', errors='ignore')
            
        except Exception as e:
            logger.error(f"Error processing TXT: {e}")
            return ""
    
    def _process_csv(self, file_content: bytes) -> str:
        """Extract text from CSV file."""
        try:
            csv_file = io.StringIO(file_content.decode('utf-8'))
            df = pd.read_csv(csv_file)
            
            # Convert DataFrame to readable text
            text_parts = [f"CSV Data with {len(df)} rows and {len(df.columns)} columns:"]
            
            # Add column headers
            text_parts.append("Columns: " + ", ".join(df.columns))
            
            # Add sample data (first 10 rows)
            for idx, row in df.head(10).iterrows():
                row_text = []
                for col, value in row.items():
                    if pd.notna(value):
                        row_text.append(f"{col}: {value}")
                text_parts.append(" | ".join(row_text))
            
            if len(df) > 10:
                text_parts.append(f"... and {len(df) - 10} more rows")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error processing CSV: {e}")
            return ""
    
    def _process_xlsx(self, file_content: bytes) -> str:
        """Extract text from XLSX file."""
        try:
            xlsx_file = io.BytesIO(file_content)
            
            # Read all sheets
            excel_data = pd.read_excel(xlsx_file, sheet_name=None)
            
            text_parts = []
            
            for sheet_name, df in excel_data.items():
                text_parts.append(f"[Sheet: {sheet_name}]")
                text_parts.append(f"Data with {len(df)} rows and {len(df.columns)} columns:")
                
                # Add column headers
                text_parts.append("Columns: " + ", ".join(df.columns))
                
                # Add sample data (first 5 rows per sheet)
                for idx, row in df.head(5).iterrows():
                    row_text = []
                    for col, value in row.items():
                        if pd.notna(value):
                            row_text.append(f"{col}: {value}")
                    text_parts.append(" | ".join(row_text))
                
                if len(df) > 5:
                    text_parts.append(f"... and {len(df) - 5} more rows")
                
                text_parts.append("")  # Empty line between sheets
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error processing XLSX: {e}")
            return ""
    
    def _process_html(self, file_content: bytes) -> str:
        """Extract text from HTML file."""
        try:
            html_text = file_content.decode('utf-8')
            
            # Use BeautifulSoup for better text extraction
            soup = BeautifulSoup(html_text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text and clean it
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            logger.error(f"Error processing HTML: {e}")
            return ""
    
    def _process_markdown(self, file_content: bytes) -> str:
        """Extract text from Markdown file."""
        try:
            md_text = file_content.decode('utf-8')
            
            # Convert markdown to HTML first, then to text
            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = False
            
            return h.handle(md_text)
            
        except Exception as e:
            logger.error(f"Error processing Markdown: {e}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and preprocess text."""
        try:
            # Remove excessive whitespace
            text = re.sub(r'\s+', ' ', text)
            
            # Remove special characters but keep basic punctuation
            text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)]', ' ', text)
            
            # Remove extra spaces
            text = text.strip()
            
            # Optional: Remove stop words and lemmatize (can be memory intensive)
            # Uncomment if you want deeper text processing
            # sentences = sent_tokenize(text)
            # cleaned_sentences = []
            # for sentence in sentences:
            #     words = word_tokenize(sentence.lower())
            #     words = [self.lemmatizer.lemmatize(word) for word in words 
            #              if word not in self.stop_words and word.isalnum()]
            #     if words:
            #         cleaned_sentences.append(' '.join(words))
            # text = ' '.join(cleaned_sentences)
            
            return text
            
        except Exception as e:
            logger.error(f"Error cleaning text: {e}")
            return text
    
    def _create_chunks(self, text: str) -> List[str]:
        """Create overlapping text chunks."""
        try:
            if len(text) <= self.chunk_size:
                return [text] if len(text) >= self.min_chunk_size else []
            
            chunks = []
            start = 0
            
            while start < len(text):
                end = start + self.chunk_size
                
                # If this is not the last chunk, try to end at a sentence boundary
                if end < len(text):
                    # Look for sentence endings within the last 200 characters
                    search_start = max(end - 200, start)
                    sentence_ends = []
                    
                    for i in range(search_start, end):
                        if text[i] in '.!?':
                            # Make sure it's followed by whitespace or end of text
                            if i + 1 < len(text) and text[i + 1].isspace():
                                sentence_ends.append(i + 1)
                    
                    if sentence_ends:
                        end = sentence_ends[-1]  # Use the last sentence end
                
                chunk = text[start:end].strip()
                
                if len(chunk) >= self.min_chunk_size:
                    chunks.append(chunk)
                
                # Move start position with overlap
                if end >= len(text):
                    break
                
                start = end - self.chunk_overlap
                
                # Ensure we make progress
                if start <= chunks[-1:] and len(chunks) > 0:
                    start = end
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error creating chunks: {e}")
            return [text] if len(text) >= self.min_chunk_size else []
    
    def _extract_metadata(self, text: str, filename: str, file_type: str) -> Dict[str, Any]:
        """Extract metadata from document."""
        try:
            metadata = {
                "filename": filename,
                "file_type": file_type,
                "text_length": len(text),
                "word_count": len(text.split()),
                "extracted_at": datetime.now().isoformat()
            }
            
            # Extract title (first meaningful line)
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            if lines:
                # Look for title-like patterns
                title_candidates = []
                for line in lines[:10]:  # Check first 10 lines
                    if 5 <= len(line) <= 100 and not line.startswith(('http', 'www')):
                        title_candidates.append(line)
                
                metadata["title"] = title_candidates[0] if title_candidates else filename
            
            # Extract dates (simple pattern matching)
            date_patterns = [
                r'\d{4}-\d{2}-\d{2}',
                r'\d{2}/\d{2}/\d{4}',
                r'[A-Za-z]+ \d{1,2}, \d{4}'
            ]
            
            dates_found = []
            for pattern in date_patterns:
                dates_found.extend(re.findall(pattern, text))
            
            if dates_found:
                metadata["dates_mentioned"] = dates_found[:5]  # First 5 dates
            
            # Extract key phrases (simple approach)
            sentences = sent_tokenize(text)
            if sentences:
                # Get first and last sentences as potential summaries
                metadata["first_sentence"] = sentences[0][:200]
                if len(sentences) > 1:
                    metadata["last_sentence"] = sentences[-1][:200]
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata: {e}")
            return {
                "filename": filename,
                "file_type": file_type,
                "error": str(e)
            }
    
    def _generate_document_hash(self, text: str) -> str:
        """Generate a hash for document deduplication."""
        return hashlib.md5(text.encode('utf-8')).hexdigest()
    
    def get_supported_types(self) -> List[str]:
        """Get list of supported file types."""
        return list(self.supported_extensions.keys())
    
    def validate_file(self, filename: str, max_size_mb: int = 50) -> Dict[str, Any]:
        """Validate file before processing."""
        try:
            file_ext = os.path.splitext(filename)[1].lower()
            
            validation = {
                "is_valid": True,
                "filename": filename,
                "file_type": file_ext,
                "errors": [],
                "warnings": []
            }
            
            # Check file extension
            if file_ext not in self.supported_extensions:
                validation["is_valid"] = False
                validation["errors"].append(f"Unsupported file type: {file_ext}")
            
            # Add other validations as needed
            # File size check would need to be done at upload time
            
            return validation
            
        except Exception as e:
            return {
                "is_valid": False,
                "filename": filename,
                "errors": [str(e)]
            }


# Global instance
_document_processor = None

def get_document_processor() -> DocumentProcessor:
    """Get global document processor instance."""
    global _document_processor
    if _document_processor is None:
        from config.settings import config
        _document_processor = DocumentProcessor(
            chunk_size=getattr(config, 'chunk_size', 1000),
            chunk_overlap=getattr(config, 'chunk_overlap', 200)
        )
    return _document_processor